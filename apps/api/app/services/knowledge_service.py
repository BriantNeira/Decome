"""Knowledge base service — documents, structured fields, customer profiles."""
from __future__ import annotations

import os
import uuid
import logging
from pathlib import Path

from fastapi import HTTPException, UploadFile
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.account_document import AccountDocument
from app.models.account_knowledge import AccountKnowledge
from app.models.customer_profile import CustomerProfile

logger = logging.getLogger(__name__)

ALLOWED_TYPES = {"pdf", "docx", "doc"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def _doc_upload_dir(account_id: uuid.UUID) -> Path:
    base = Path(settings.upload_dir) / "accounts" / str(account_id) / "docs"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _extract_text(file_path: Path, file_type: str) -> str:
    """Extract plain text from PDF or DOCX."""
    text = ""
    try:
        if file_type == "pdf":
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif file_type in ("docx", "doc"):
            import docx
            doc = docx.Document(str(file_path))
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", file_path, exc)
    return text.strip()


async def upload_document(
    db: AsyncSession,
    account_id: uuid.UUID,
    file: UploadFile,
    uploaded_by: uuid.UUID,
) -> AccountDocument:
    # Validate extension
    original_name = file.filename or "file"
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=422, detail=f"File type '.{ext}' not allowed. Use PDF, DOCX or DOC.")

    # Read content
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=422, detail="File exceeds 10 MB limit.")

    # Save to disk
    file_id = uuid.uuid4()
    safe_name = f"{file_id}_{original_name}"
    dest = _doc_upload_dir(account_id) / safe_name
    dest.write_bytes(content)

    # Extract text
    content_text = _extract_text(dest, ext)

    # Save DB record
    doc = AccountDocument(
        account_id=account_id,
        filename=original_name,
        file_path=str(dest),
        file_type=ext,
        content_text=content_text or None,
        uploaded_by=uploaded_by,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


async def list_documents(db: AsyncSession, account_id: uuid.UUID) -> list[AccountDocument]:
    result = await db.execute(
        select(AccountDocument)
        .where(AccountDocument.account_id == account_id)
        .order_by(AccountDocument.uploaded_at.desc())
    )
    return list(result.scalars().all())


async def delete_document(db: AsyncSession, doc_id: uuid.UUID) -> None:
    result = await db.execute(select(AccountDocument).where(AccountDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    # Delete file
    try:
        Path(doc.file_path).unlink(missing_ok=True)
    except Exception as exc:
        logger.warning("Could not delete file %s: %s", doc.file_path, exc)
    await db.delete(doc)
    await db.commit()


async def get_knowledge(db: AsyncSession, account_id: uuid.UUID) -> AccountKnowledge | None:
    result = await db.execute(
        select(AccountKnowledge).where(AccountKnowledge.account_id == account_id)
    )
    return result.scalar_one_or_none()


async def upsert_knowledge(
    db: AsyncSession,
    account_id: uuid.UUID,
    updated_by: uuid.UUID,
    **fields,
) -> AccountKnowledge:
    obj = await get_knowledge(db, account_id)
    if obj is None:
        obj = AccountKnowledge(account_id=account_id, updated_by=updated_by, **fields)
        db.add(obj)
    else:
        for k, v in fields.items():
            setattr(obj, k, v)
        obj.updated_by = updated_by
    await db.commit()
    await db.refresh(obj)
    return obj


async def get_latest_profile(db: AsyncSession, account_id: uuid.UUID) -> CustomerProfile | None:
    result = await db.execute(
        select(CustomerProfile)
        .where(CustomerProfile.account_id == account_id)
        .order_by(CustomerProfile.version.desc())
        .limit(1)
    )
    return result.scalar_one_or_none()


async def list_profile_history(db: AsyncSession, account_id: uuid.UUID) -> list[CustomerProfile]:
    result = await db.execute(
        select(CustomerProfile)
        .where(CustomerProfile.account_id == account_id)
        .order_by(CustomerProfile.version.desc())
    )
    return list(result.scalars().all())


async def build_profile(
    db: AsyncSession,
    account_id: uuid.UUID,
    generated_by: uuid.UUID,
) -> CustomerProfile:
    """Build or rebuild customer profile using LLM."""
    from app.services import llm_service, budget_service
    from app.models.account_note import AccountNote
    from app.models.account import Account

    # Load LLM config
    config = await llm_service.get_llm_config(db)
    if config is None or not config.is_active:
        raise HTTPException(status_code=422, detail="LLM is not configured or not active.")
    if not config.api_key:
        raise HTTPException(status_code=422, detail="LLM API key is not set.")

    # Budget check for the requesting user
    await budget_service.check_budget(db, generated_by)

    # Load account name
    account_result = await db.execute(select(Account).where(Account.id == account_id))
    account = account_result.scalar_one_or_none()
    if account is None:
        raise HTTPException(status_code=404, detail="Account not found")

    # Load knowledge
    knowledge_obj = await get_knowledge(db, account_id)
    knowledge_dict = {}
    if knowledge_obj:
        knowledge_dict = {
            "website": knowledge_obj.website or "N/A",
            "main_email": knowledge_obj.main_email or "N/A",
            "industry": knowledge_obj.industry or "N/A",
            "account_type": knowledge_obj.account_type or "N/A",
            "observations": knowledge_obj.observations or "N/A",
        }

    # Load documents text
    docs = await list_documents(db, account_id)
    docs_text = "\n\n".join(
        d.content_text for d in docs if d.content_text
    )[:3000]

    # Load last 5 notes
    notes_result = await db.execute(
        select(AccountNote)
        .where(AccountNote.account_id == account_id)
        .order_by(AccountNote.created_at.desc())
        .limit(5)
    )
    notes = notes_result.scalars().all()
    notes_text = "\n".join(n.content for n in notes if n.content)

    # Call LLM
    profile_text, tokens_used = await llm_service.generate_profile(
        config,
        account_name=account.name,
        knowledge=knowledge_dict,
        docs_text=docs_text,
        notes_text=notes_text,
    )

    # Next version number
    latest = await get_latest_profile(db, account_id)
    next_version = (latest.version + 1) if latest else 1

    # Save profile
    profile = CustomerProfile(
        account_id=account_id,
        profile_text=profile_text,
        version=next_version,
        generated_by=generated_by,
        tokens_used=tokens_used,
    )
    db.add(profile)
    await db.commit()
    await db.refresh(profile)

    # Log token usage
    await budget_service.log_usage(
        db,
        user_id=generated_by,
        account_id=account_id,
        reminder_id=None,
        tokens=tokens_used,
        provider=config.provider,
        model=config.model,
    )

    return profile
